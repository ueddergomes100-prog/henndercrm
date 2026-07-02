from __future__ import annotations

import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
OUTPUT_DOCX = DOCS_DIR / "Hennder_CRM_Apresentacao_Funcionalidades.docx"

DOCUMENTS_PLUGIN_DIR = Path(
    r"C:\Users\uedde\.codex\plugins\cache\openai-primary-runtime\documents"
)
DOCUMENTS_PLUGIN_VERSIONS = sorted(
    (path for path in DOCUMENTS_PLUGIN_DIR.glob("*") if path.is_dir()),
    key=lambda path: path.name,
)
SKILL_SCRIPTS = (
    DOCUMENTS_PLUGIN_VERSIONS[-1] / "skills" / "documents" / "scripts"
    if DOCUMENTS_PLUGIN_VERSIONS
    else DOCUMENTS_PLUGIN_DIR / "skills" / "documents" / "scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


# compact_reference_guide preset with named Hennder brochure overrides:
# Arial typography, branded colors, full-bleed cover, and 10.5 pt body copy.
NAVY = "082F63"
BLUE = "0753A6"
CYAN = "0B9FB0"
TEAL = "14BD8A"
ORANGE = "EA580C"
AMBER = "F59E0B"
PURPLE = "5B35C9"
DARK = "02040A"
INK = "123252"
SLATE = "475569"
MUTED = "64748B"
PALE_BLUE = "EAF3FB"
PALE_CYAN = "E7F8FA"
PALE_ORANGE = "FFF3E8"
PALE_GREEN = "EAFBF4"
PALE_PURPLE = "F2EEFF"
WHITE = "FFFFFF"
LIGHT_BORDER = "D8E6F2"
CONTENT_WIDTH_DXA = 9360


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def font_path(bold: bool = False) -> str:
    candidate = (
        Path(r"C:\Windows\Fonts\segoeuib.ttf")
        if bold
        else Path(r"C:\Windows\Fonts\segoeui.ttf")
    )
    if candidate.exists():
        return str(candidate)
    return str(
        Path(r"C:\Windows\Fonts\arialbd.ttf")
        if bold
        else Path(r"C:\Windows\Fonts\arial.ttf")
    )


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold), size)


def lerp(a: int, b: int, value: float) -> int:
    return int(a + (b - a) * value)


def hex_tuple(value: str) -> tuple[int, int, int]:
    return tuple(int(value[index : index + 2], 16) for index in (0, 2, 4))


def draw_gradient(
    image: Image.Image,
    start: str,
    end: str,
    vertical: bool = True,
) -> None:
    draw = ImageDraw.Draw(image)
    start_rgb = hex_tuple(start)
    end_rgb = hex_tuple(end)
    length = image.height if vertical else image.width
    for position in range(length):
        ratio = position / max(length - 1, 1)
        color = tuple(lerp(start_rgb[i], end_rgb[i], ratio) for i in range(3))
        if vertical:
            draw.line((0, position, image.width, position), fill=color)
        else:
            draw.line((position, 0, position, image.height), fill=color)


def rounded_rect(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    radius: int = 28,
    outline: str | None = None,
    width: int = 1,
) -> None:
    draw.rounded_rectangle(
        box,
        radius=radius,
        fill=f"#{fill}",
        outline=f"#{outline}" if outline else None,
        width=width,
    )


def create_cover() -> Path:
    path = ASSETS_DIR / "cover.png"
    image = Image.new("RGB", (2550, 3300), f"#{DARK}")
    draw_gradient(image, DARK, NAVY)
    draw = ImageDraw.Draw(image, "RGBA")

    for x, y, radius, color, alpha in [
        (2180, 420, 620, CYAN, 55),
        (2100, 2850, 850, TEAL, 35),
        (180, 2920, 520, PURPLE, 48),
        (400, 320, 310, BLUE, 80),
    ]:
        draw.ellipse(
            (x - radius, y - radius, x + radius, y + radius),
            fill=(*hex_tuple(color), alpha),
        )

    for offset in range(-800, 2800, 280):
        draw.line(
            (offset, 2520, offset + 1200, 3300),
            fill=(255, 255, 255, 12),
            width=4,
        )

    rounded_rect(draw, (220, 220, 760, 340), CYAN, 60)
    draw.text(
        (490, 280),
        "APRESENTAÇÃO FUNCIONAL",
        font=pil_font(40, True),
        fill=f"#{WHITE}",
        anchor="mm",
    )

    rounded_rect(draw, (220, 640, 470, 890), WHITE, 58)
    draw.text(
        (345, 765),
        "H",
        font=pil_font(145, True),
        fill=f"#{BLUE}",
        anchor="mm",
    )
    draw.text(
        (220, 1080),
        "Hennder CRM",
        font=pil_font(178, True),
        fill=f"#{WHITE}",
    )
    draw.text(
        (228, 1315),
        "Inteligência Comercial e Recompra",
        font=pil_font(66),
        fill="#D9F8FF",
    )

    draw.line((220, 1485, 1630, 1485), fill=f"#{CYAN}", width=12)
    title = "Dados que viram\nrelacionamento, recorrência\ne novas vendas."
    draw.multiline_text(
        (220, 1620),
        title,
        font=pil_font(100, True),
        fill=f"#{WHITE}",
        spacing=24,
    )

    rounded_rect(draw, (220, 2625, 2330, 2970), "07111F", 42, "2D597C", 3)
    draw.text(
        (310, 2720),
        "Uma plataforma para organizar clientes, recuperar oportunidades,",
        font=pil_font(43),
        fill="#DCEEFF",
    )
    draw.text(
        (310, 2805),
        "antecipar recompras e orientar a equipe comercial com contexto.",
        font=pil_font(43),
        fill="#DCEEFF",
    )
    draw.text(
        (220, 3150),
        "JUNHO DE 2026",
        font=pil_font(34, True),
        fill="#A9D8F7",
    )
    image.save(path, quality=96)
    return path


def create_integration_flow() -> Path:
    path = ASSETS_DIR / "integration-flow.png"
    image = Image.new("RGB", (1800, 620), f"#{PALE_BLUE}")
    draw = ImageDraw.Draw(image)
    boxes = [
        ("ERP DO CLIENTE", "Clientes, produtos,\nvendas e vendedores", BLUE),
        ("CONEXÃO ADAPTÁVEL", "Importação, API ou\nsincronização agendada", CYAN),
        ("HENNDER CRM", "Regras, indicadores\ne inteligência comercial", PURPLE),
        ("EQUIPE COMERCIAL", "Contatos, agenda,\noportunidades e resultados", TEAL),
    ]
    x_positions = [40, 480, 920, 1360]
    for index, ((title, subtitle, color), x) in enumerate(zip(boxes, x_positions)):
        rounded_rect(draw, (x, 90, x + 390, 525), WHITE, 30, LIGHT_BORDER, 3)
        rounded_rect(draw, (x + 28, 120, x + 362, 215), color, 22)
        draw.text(
            (x + 195, 168),
            title,
            font=pil_font(28, True),
            fill=f"#{WHITE}",
            anchor="mm",
        )
        draw.multiline_text(
            (x + 195, 320),
            subtitle,
            font=pil_font(30),
            fill=f"#{INK}",
            anchor="mm",
            align="center",
            spacing=12,
        )
        if index < len(boxes) - 1:
            start = x + 395
            end = x_positions[index + 1] - 10
            draw.line((start, 310, end, 310), fill=f"#{CYAN}", width=10)
            draw.polygon(
                [(end - 18, 290), (end + 8, 310), (end - 18, 330)],
                fill=f"#{CYAN}",
            )
    image.save(path, quality=95)
    return path


def create_dashboard_preview() -> Path:
    path = ASSETS_DIR / "dashboard-preview.png"
    image = Image.new("RGB", (1800, 850), f"#{DARK}")
    draw = ImageDraw.Draw(image)
    draw_gradient(image, "050914", "071D35")
    draw = ImageDraw.Draw(image, "RGBA")

    draw.text(
        (85, 60),
        "Visão executiva",
        font=pil_font(48, True),
        fill=f"#{WHITE}",
    )
    draw.text(
        (85, 125),
        "Indicadores comerciais em uma única leitura",
        font=pil_font(28),
        fill="#A9C8E5",
    )

    kpis = [
        ("CLIENTES ATIVOS", "1.284", PURPLE),
        ("ALERTAS DE RECOMPRA", "86", BLUE),
        ("POTENCIAL RECUPERÁVEL", "R$ 74,8 mil", ORANGE),
        ("OPORTUNIDADES", "42", TEAL),
    ]
    for index, (label, value, color) in enumerate(kpis):
        x = 85 + index * 415
        rounded_rect(draw, (x, 215, x + 365, 410), color, 28)
        draw.text((x + 26, 252), label, font=pil_font(20, True), fill="#EAF8FF")
        draw.text((x + 26, 305), value, font=pil_font(48, True), fill=f"#{WHITE}")
        draw.line((x + 26, 382, x + 330, 382), fill=(255, 255, 255, 55), width=3)

    rounded_rect(draw, (85, 465, 1080, 775), "0A1423", 28, "203047", 3)
    draw.text(
        (125, 505),
        "Evolução de recompra",
        font=pil_font(28, True),
        fill=f"#{WHITE}",
    )
    points = [(160, 710), (300, 675), (445, 690), (590, 610), (735, 630), (880, 545), (1025, 565)]
    for start, end in zip(points, points[1:]):
        draw.line((*start, *end), fill=f"#{CYAN}", width=10)
    for point in points:
        draw.ellipse(
            (point[0] - 10, point[1] - 10, point[0] + 10, point[1] + 10),
            fill=f"#{TEAL}",
        )

    rounded_rect(draw, (1120, 465, 1715, 775), "0A1423", 28, "203047", 3)
    draw.text(
        (1160, 505),
        "Prioridades do dia",
        font=pil_font(28, True),
        fill=f"#{WHITE}",
    )
    priorities = [
        ("Recompra prevista", 0.84, CYAN),
        ("Cliente em risco", 0.68, ORANGE),
        ("Venda cruzada", 0.52, PURPLE),
    ]
    for index, (label, ratio, color) in enumerate(priorities):
        y = 575 + index * 70
        draw.text((1160, y), label, font=pil_font(22), fill="#C8D9E8")
        draw.rounded_rectangle((1440, y + 5, 1670, y + 27), radius=11, fill="#17263A")
        draw.rounded_rectangle(
            (1440, y + 5, 1440 + int(230 * ratio), y + 27),
            radius=11,
            fill=f"#{color}",
        )
    image.save(path, quality=95)
    return path


def create_commercial_journey() -> Path:
    path = ASSETS_DIR / "commercial-journey.png"
    image = Image.new("RGB", (1800, 660), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    stages = [
        ("1", "ENTENDER", "Dashboard e\nsegmentação", BLUE),
        ("2", "PRIORIZAR", "Risco, recompra\ne potencial", ORANGE),
        ("3", "AGIR", "Contato, agenda\ne oportunidade", CYAN),
        ("4", "ACOMPANHAR", "Histórico, status\ne carteira", PURPLE),
        ("5", "MEDIR", "Relatórios e\nresultados", TEAL),
    ]
    for index, (number, title, detail, color) in enumerate(stages):
        center_x = 180 + index * 360
        if index < len(stages) - 1:
            draw.line(
                (center_x + 105, 210, center_x + 255, 210),
                fill=f"#{LIGHT_BORDER}",
                width=14,
            )
        draw.ellipse(
            (center_x - 82, 128, center_x + 82, 292),
            fill=f"#{color}",
        )
        draw.text(
            (center_x, 210),
            number,
            font=pil_font(58, True),
            fill=f"#{WHITE}",
            anchor="mm",
        )
        draw.text(
            (center_x, 355),
            title,
            font=pil_font(26, True),
            fill=f"#{INK}",
            anchor="mm",
        )
        draw.multiline_text(
            (center_x, 455),
            detail,
            font=pil_font(27),
            fill=f"#{SLATE}",
            anchor="mm",
            align="center",
            spacing=8,
        )
    image.save(path, quality=95)
    return path


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str = LIGHT_BORDER, size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = tc_borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            tc_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Arial"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    heading_1.font.size = Pt(18)
    heading_1.font.bold = True
    heading_1.font.color.rgb = rgb(NAVY)
    heading_1.paragraph_format.space_before = Pt(15)
    heading_1.paragraph_format.space_after = Pt(7)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    heading_2.font.size = Pt(12.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = rgb(BLUE)
    heading_2.paragraph_format.space_before = Pt(9)
    heading_2.paragraph_format.space_after = Pt(3)
    heading_2.paragraph_format.keep_with_next = True


def configure_body_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("HENNDER CRM  |  INTELIGÊNCIA COMERCIAL E RECOMPRA")
    set_run_font(run, size=8.3, color=BLUE, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    run = footer_p.add_run("Apresentação funcional  •  Junho de 2026  •  ")
    set_run_font(run, size=8.3, color=MUTED)
    add_page_field(footer_p)


def add_section_banner(
    document: Document,
    number: str,
    kicker: str,
    title: str,
    subtitle: str,
    color: str = BLUE,
) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_border(cell, color, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=0,
        cell_margins_dxa={"top": 150, "bottom": 150, "start": 210, "end": 210},
    )
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{number}  {kicker.upper()}")
    set_run_font(run, size=8.5, color="D9F8FF", bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=20, color=WHITE, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=10.3, color="EAF8FF")
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.2


def add_lead(document: Document, text: str, color: str = CYAN) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    run = paragraph.add_run(text)
    set_run_font(run, size=11.2, color=INK, bold=True)


def add_label_detail_table(
    document: Document,
    rows: list[tuple[str, str]],
    *,
    label_color: str = BLUE,
    alternate: bool = True,
) -> None:
    table = document.add_table(rows=1, cols=2)
    header = table.rows[0]
    header.cells[0].text = "RECURSO"
    header.cells[1].text = "O QUE ENTREGA"
    set_repeat_table_header(header)
    for cell in header.cells:
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY, 8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=8.8, color=WHITE, bold=True)

    for index, (label, detail) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        fill = PALE_BLUE if alternate and index % 2 == 0 else WHITE
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_border(cell, LIGHT_BORDER, 7)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, size=9.3, color=INK)
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.3, color=label_color, bold=True)

    apply_table_geometry(
        table,
        [2500, 6860],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 105, "bottom": 105, "start": 135, "end": 135},
    )
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_highlight_strip(
    document: Document,
    items: list[tuple[str, str, str]],
) -> None:
    table = document.add_table(rows=1, cols=len(items))
    widths = [CONTENT_WIDTH_DXA // len(items)] * len(items)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    for cell, (value, label, color) in zip(table.rows[0].cells, items):
        set_cell_shading(cell, color)
        set_cell_border(cell, WHITE, 12)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        set_run_font(run, size=17, color=WHITE, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=8.3, color=WHITE, bold=True)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=0,
        cell_margins_dxa={"top": 125, "bottom": 125, "start": 90, "end": 90},
    )


def add_note_box(
    document: Document,
    title: str,
    body: str,
    *,
    fill: str = PALE_CYAN,
    accent: str = CYAN,
) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, 10)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=10.2, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(body)
    set_run_font(run, size=9.4, color=INK)
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 125, "bottom": 125, "start": 160, "end": 160},
    )


def add_page_break(document: Document) -> None:
    document.add_page_break()


def build_document() -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    cover = create_cover()
    flow = create_integration_flow()
    dashboard = create_dashboard_preview()
    journey = create_commercial_journey()

    document = Document()
    configure_styles(document)

    cover_section = document.sections[0]
    cover_section.page_width = Inches(8.5)
    cover_section.page_height = Inches(11)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)
    cover_section.different_first_page_header_footer = True
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(cover), width=Inches(8.5), height=Inches(11))

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_body_section(body_section)

    # Page 2
    add_section_banner(
        document,
        "01",
        "Visão do produto",
        "Um CRM criado para transformar dados em ação comercial",
        "A equipe sabe quem priorizar, por que entrar em contato e qual é o próximo passo.",
        BLUE,
    )
    add_lead(
        document,
        "O Hennder CRM organiza o relacionamento com clientes e usa o histórico comercial "
        "para revelar risco, potencial de recompra e oportunidades de venda.",
    )
    add_highlight_strip(
        document,
        [
            ("RECUPERAR", "clientes sem compra", ORANGE),
            ("ANTECIPAR", "o momento da recompra", CYAN),
            ("CRESCER", "com venda cruzada", PURPLE),
            ("ORGANIZAR", "a rotina da equipe", TEAL),
        ],
    )
    document.add_heading("Integração independente do ERP", level=2)
    paragraph = document.add_paragraph(
        "A plataforma trabalha com uma camada de conexão adaptável. Cada implantação pode "
        "receber dados por consulta somente leitura, API, arquivos ou sincronização agendada, "
        "de acordo com a estrutura e as permissões do ERP utilizado pelo cliente."
    )
    paragraph.paragraph_format.space_after = Pt(7)
    document.add_picture(str(flow), width=Inches(6.5))
    add_note_box(
        document,
        "Base comum de dados",
        "Clientes, produtos, vendas, itens vendidos, vendedores, datas, valores, filiais e "
        "status comerciais formam a base necessária para ativar os indicadores.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    # Page 3
    add_page_break(document)
    add_section_banner(
        document,
        "02",
        "Gestão executiva",
        "Dashboard comercial inteligente",
        "Os principais sinais da operação aparecem de forma visual, objetiva e priorizada.",
        PURPLE,
    )
    add_lead(
        document,
        "Uma leitura rápida da carteira, da recorrência e das oportunidades ajuda gestores e "
        "vendedores a começarem o dia pelo que realmente merece atenção.",
        PURPLE,
    )
    document.add_picture(str(dashboard), width=Inches(6.5))
    add_label_detail_table(
        document,
        [
            ("KPIs comerciais", "Indicadores de clientes, recompra, receita, risco e potencial recuperável."),
            ("Gráficos de evolução", "Acompanhamento visual da recorrência e do desempenho por categoria."),
            ("Ranking de prioridades", "Clientes e vendedores ordenados pelos sinais mais relevantes."),
            ("Prévia de inatividade", "Resumo dos clientes sem compra e acesso direto à Central de Recuperação."),
        ],
        label_color=PURPLE,
    )

    # Page 4
    add_page_break(document)
    add_section_banner(
        document,
        "03",
        "Conhecimento do cliente",
        "Carteira segmentada e Perfil 360°",
        "Cada cliente é visto pelo cadastro, comportamento de compra, relacionamento e potencial.",
        BLUE,
    )
    add_label_detail_table(
        document,
        [
            ("Consulta de clientes", "Busca por nome ou cidade e filtros por vendedor, status e qualidade cadastral."),
            ("Perfil 360°", "Dados cadastrais, compras, itens, alertas, oportunidades e histórico de contatos."),
            ("Indicadores individuais", "Ticket médio, frequência, dias sem compra, score e potencial perdido."),
            ("Vendedor preferencial", "Identificação do profissional com maior relacionamento no histórico de compras."),
            ("Qualidade cadastral", "Pontuação para telefone, WhatsApp, e-mail, documentos, endereço e localização."),
            ("Status de inatividade", "Classificação entre ativo, atenção, risco e perdido conforme o tempo sem compra."),
        ],
        label_color=BLUE,
    )
    document.add_heading("Saúde da base", level=2)
    paragraph = document.add_paragraph(
        "O módulo de qualidade mostra cadastros sem WhatsApp, telefone, cidade ou CPF/CNPJ, "
        "calcula o score médio e destaca os registros que precisam de correção."
    )
    paragraph.paragraph_format.space_after = Pt(7)
    add_highlight_strip(
        document,
        [
            ("CADASTRO", "qualidade mensurável", BLUE),
            ("HISTÓRICO", "compras e contatos", CYAN),
            ("AFINIDADE", "vendedor recomendado", PURPLE),
            ("POTENCIAL", "valor a recuperar", ORANGE),
        ],
    )
    add_note_box(
        document,
        "Visão prática",
        "Além de entender quem é o cliente, a equipe enxerga o contexto necessário para fazer "
        "uma abordagem mais relevante e menos genérica.",
        fill=PALE_PURPLE,
        accent=PURPLE,
    )

    # Page 5
    add_page_break(document)
    add_section_banner(
        document,
        "04",
        "Recorrência e relacionamento",
        "Recuperação de clientes e alertas de recompra",
        "O sistema transforma ausência de compra e ciclos de consumo em filas de ação.",
        ORANGE,
    )
    add_label_detail_table(
        document,
        [
            ("Central de Recuperação", "Clientes sem compra ordenados por urgência, dias de inatividade e potencial."),
            ("Registro de retorno", "Resultado, observação, canal, responsável e próxima data de contato."),
            ("Histórico de contatos", "Tentativas registradas permanecem disponíveis para continuidade do atendimento."),
            ("Alertas de recompra", "Previsões por produto, departamento, palavra-chave e histórico individual."),
            ("Priorização operacional", "Filtros de hoje, próximos dias, atrasados e nível de prioridade."),
            ("Status persistentes", "Ações para marcar alertas como contatados, convertidos ou ignorados."),
        ],
        label_color=ORANGE,
    )
    add_note_box(
        document,
        "Regra essencial de recuperação",
        "Uma tentativa de contato não elimina o alerta. A oportunidade permanece ativa até "
        "que exista uma nova compra ou uma decisão comercial registrada.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    document.add_heading("WhatsApp com contexto", level=2)
    paragraph = document.add_paragraph(
        "Os atalhos abrem uma conversa com número normalizado e mensagem relacionada ao "
        "cliente, ao produto ou à oportunidade selecionada. A automação oficial pode ser "
        "adicionada posteriormente conforme consentimento, templates e regras da operação."
    )
    paragraph.paragraph_format.space_after = Pt(0)

    # Page 6
    add_page_break(document)
    add_section_banner(
        document,
        "05",
        "Execução comercial",
        "Carteira, oportunidades e agenda em um só fluxo",
        "A equipe recebe contexto para vender e ferramentas para organizar a próxima ação.",
        TEAL,
    )
    add_label_detail_table(
        document,
        [
            ("Carteira do vendedor", "Clientes preferenciais, clientes em risco, alertas, potencial e conversão."),
            ("Oportunidades", "Sugestões de venda cruzada com cliente, produto de origem, recomendação e confiança."),
            ("Gestão de oportunidades", "Criação, edição, responsável, evolução de status e exclusão."),
            ("Agenda semanal", "Ligações, visitas, retornos e recompras previstas com visão operacional."),
            ("Gestão de compromissos", "Criação, edição e exclusão de eventos, com controle por vendedor."),
            ("Acesso por perfil", "Administrador e supervisor veem a operação; vendedor atua na própria carteira."),
        ],
        label_color=TEAL,
    )
    add_highlight_strip(
        document,
        [
            ("FOCO", "prioridade por vendedor", TEAL),
            ("CONTEXTO", "abordagem personalizada", BLUE),
            ("CADÊNCIA", "agenda de retornos", CYAN),
            ("CONTROLE", "status e responsável", PURPLE),
        ],
    )
    add_note_box(
        document,
        "Continuidade da operação",
        "Contatos, status de alertas, oportunidades e eventos da agenda ficam persistidos, "
        "evitando que a rotina dependa de anotações dispersas.",
        fill=PALE_GREEN,
        accent=TEAL,
    )

    # Page 7
    add_page_break(document)
    add_section_banner(
        document,
        "06",
        "Inteligência e gestão",
        "Análises, relatórios, segurança e experiência de uso",
        "Recursos de apoio ajudam a interpretar a base, acompanhar resultados e governar o acesso.",
        CYAN,
    )
    add_label_detail_table(
        document,
        [
            ("IA Comercial", "Interface de perguntas e respostas orientada pelos dados atuais do CRM."),
            ("Análises disponíveis", "Risco, potencial, vendedores, produtos, cadastro e prioridades de contato."),
            ("Relatórios", "Indicadores de perda, recuperação, recorrência e faturamento recuperado."),
            ("Perfis de acesso", "Experiências separadas para administrador, supervisor e vendedor."),
            ("Sessão protegida", "Autenticação com cookie HTTP-only e escopo de dados conforme o perfil."),
            ("Temas visuais", "Seleção entre modo claro e dark profundo, com preferência salva pelo usuário."),
            ("Persistência operacional", "Informações da rotina comercial mantidas no banco remoto."),
        ],
        label_color=CYAN,
    )
    add_note_box(
        document,
        "Evolução planejada da inteligência",
        "O módulo analítico atual trabalha com regras e dados disponíveis no CRM. A arquitetura "
        "permite incorporar modelos externos quando dados, governança e objetivos estiverem maduros.",
        fill=PALE_CYAN,
        accent=CYAN,
    )
    add_highlight_strip(
        document,
        [
            ("CLARO", "leitura leve", BLUE),
            ("DARK", "contraste profundo", DARK),
            ("PERFIS", "acesso controlado", PURPLE),
            ("DADOS", "rotina persistente", TEAL),
        ],
    )

    # Page 8
    add_page_break(document)
    add_section_banner(
        document,
        "07",
        "Jornada completa",
        "Da leitura dos dados à ação e ao resultado",
        "O Hennder CRM conecta gestão, relacionamento e execução em uma rotina comercial contínua.",
        NAVY,
    )
    document.add_picture(str(journey), width=Inches(6.5))
    add_label_detail_table(
        document,
        [
            ("Dashboard", "Visão geral e prioridades."),
            ("Clientes e Perfil 360°", "Contexto completo para relacionamento."),
            ("Recuperação e Recompra", "Ações para trazer clientes de volta."),
            ("Carteira e Oportunidades", "Foco comercial por vendedor."),
            ("Agenda", "Organização da cadência de contatos."),
            ("IA e Relatórios", "Interpretação e acompanhamento dos resultados."),
        ],
        label_color=NAVY,
    )
    add_note_box(
        document,
        "Compatível com o ecossistema do cliente",
        "A integração é desenhada sobre os dados disponíveis e as regras do negócio, permitindo "
        "que o Hennder CRM trabalhe com diferentes ERPs e cenários de implantação.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(11)
    closing.paragraph_format.space_after = Pt(3)
    run = closing.add_run("Mais relacionamento. Mais recorrência. Mais venda com contexto.")
    set_run_font(run, size=15.5, color=NAVY, bold=True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("HENNDER CRM  •  INTELIGÊNCIA COMERCIAL E RECOMPRA")
    set_run_font(run, size=9, color=CYAN, bold=True)

    document.core_properties.title = "Hennder CRM - Apresentação de Funcionalidades"
    document.core_properties.subject = "Visão comercial e funcional do Hennder CRM"
    document.core_properties.author = "Hennder CRM"
    document.core_properties.keywords = (
        "CRM, inteligência comercial, recompra, recuperação de clientes, vendas"
    )
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_document())
